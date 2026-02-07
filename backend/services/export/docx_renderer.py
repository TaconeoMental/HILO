import os
import io

import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.export.script_parser import parse_script


STRONG_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_runs(paragraph, text):
    text = text or ""
    last = 0
    for match in STRONG_RE.finditer(text):
        start, end = match.span()
        if start > last:
            paragraph.add_run(text[last:start])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.font.bold = True
        last = end
    if last < len(text):
        paragraph.add_run(text[last:])


def _add_vida_table(doc, blocks, project_dir):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.rows[0].cells[0]

    title_p = cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("Vida interna")
    title_run.font.bold = True
    title_run.font.all_caps = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(64, 64, 64)

    for block in blocks or []:
        b_type = block.get("type")
        if b_type == "paragraph":
            lines = (block.get("text", "") or "").split("\n") or [""]
            for line in lines:
                para = cell.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                _add_markdown_runs(para, line)
        elif b_type == "image":
            img_rel = os.path.basename(block.get("src", "") or "")
            if not img_rel:
                continue
            img_path = os.path.join(project_dir, "photos", img_rel)
            if os.path.exists(img_path):
                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    para.add_run().add_picture(img_path, width=Inches(4))
                except Exception:
                    para.add_run(block.get("alt", "Foto"))
            else:
                para = cell.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.add_run("[Imagen no encontrada: %s]" % img_rel)

    if cell.paragraphs:
        cell.paragraphs[-1].paragraph_format.space_after = Pt(0)


def render_docx_bytes(content, project_id, project_dir):
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(12)
    nodes = parse_script(content)
    prev_blank = False

    for node in nodes:
        n_type = node.get("type")

        if n_type == "blank":
            if prev_blank:
                continue
            doc.add_paragraph("")
            prev_blank = True
            continue

        prev_blank = False

        if n_type == "heading":
            level = node.get("level", 1)
            para = doc.add_heading("", level=level)
            _add_markdown_runs(para, node.get("text", ""))
            para.paragraph_format.space_after = Pt(0)
            continue

        if n_type == "hr":
            para = doc.add_paragraph("—")
            para.paragraph_format.space_after = Pt(0)
            continue

        if n_type == "paragraph":
            para = doc.add_paragraph()
            _add_markdown_runs(para, node.get("text", ""))
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue

        if n_type == "image":
            img_rel = os.path.basename(node.get("src", "") or "")
            img_path = os.path.join(project_dir, "photos", img_rel)
            if os.path.exists(img_path):
                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.add_run().add_picture(img_path, width=Inches(4))
                except Exception:
                    doc.add_paragraph(node.get("alt", "Foto"))
            else:
                doc.add_paragraph("[Imagen no encontrada: %s]" % img_rel)
            continue

        if n_type == "vida":
            _add_vida_table(doc, node.get("blocks", []), project_dir)
            continue

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
